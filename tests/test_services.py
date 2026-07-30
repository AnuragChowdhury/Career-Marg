"""
Automated Service Unit Tests for Career मार्ग.
Uses Python built-in unittest framework for zero-dependency test runner.
Tests core service logic, scoring algorithms, schemas, and database operations.
"""

import unittest
import os
import json

from models.schemas import (
    CandidateProfile, PersonalInfo, DocumentAnalysis,
    Education, WorkExperience, Project, ATSResult, FactorScores
)
from utils.helpers import validate_uploaded_file, clean_text, extract_skills_from_text
from utils.scoring import calculate_ats_score, compute_weighted_ats_overall, calculate_industry_readiness
from services.mistral_ocr_service import MistralOCRService
from services.document_service import DocumentService
from services.resume_parser import ResumeParser
from services.ats_service import ATSService
from services.skill_gap_service import SkillGapService
from services.interview_service import InterviewService
from services.career_service import CareerService
from services.recommendation_service import RecommendationService
from services.profile_service import ProfileService
from data.database import save_candidate_profile, init_db


class TestCareerMargServices(unittest.TestCase):

    def test_file_validation(self):
        valid, msg = validate_uploaded_file("resume.pdf", 1024 * 1024)
        self.assertTrue(valid)
        
        valid_img, _ = validate_uploaded_file("resume.png", 500 * 1024)
        self.assertTrue(valid_img)

        valid_docx, _ = validate_uploaded_file("resume.docx", 1024 * 1024)
        self.assertTrue(valid_docx)

        valid_txt, _ = validate_uploaded_file("resume.txt", 1024)
        self.assertTrue(valid_txt)

        invalid_ext, err = validate_uploaded_file("resume.exe", 1024)
        self.assertFalse(invalid_ext)
        self.assertIn("Unsupported file format", err)

    def test_docx_and_txt_extraction(self):
        ocr_service = MistralOCRService()
        
        # Test TXT extraction
        txt_bytes = b"John Doe\nSoftware Engineer\nSkills: Python, Streamlit, SQL"
        txt_res = ocr_service.process_document(txt_bytes, "sample_resume.txt")
        self.assertTrue(txt_res.get("success"))
        self.assertIn("John Doe", txt_res.get("text"))
        self.assertEqual(txt_res.get("engine"), "Plain Text Parser")

        # Test DOCX extraction using python-docx Document
        import docx
        import io
        doc = docx.Document()
        doc.add_paragraph("Jane Smith - Data Scientist")
        doc.add_paragraph("Skills: Python, Machine Learning, PyTorch")
        bio = io.BytesIO()
        doc.save(bio)
        docx_bytes = bio.getvalue()

        docx_res = ocr_service.process_document(docx_bytes, "sample_resume.docx")
        self.assertTrue(docx_res.get("success"))
        self.assertIn("Jane Smith", docx_res.get("text"))
        self.assertEqual(docx_res.get("engine"), "DOCX Document Parser")

    def test_text_cleaning_and_skill_extraction(self):
        raw_sample = "Developed   a  machine learning system\n\nusing Python, PyTorch, and Docker."
        cleaned = clean_text(raw_sample)
        self.assertIn("Developed a machine learning system", cleaned)
        
        skills = extract_skills_from_text(cleaned, ["Python", "PyTorch", "Docker", "Java"])
        self.assertIn("Python", skills)
        self.assertIn("PyTorch", skills)
        self.assertIn("Docker", skills)
        self.assertNotIn("Java", skills)

    def test_ats_scoring_formula(self):
        factors = calculate_ats_score(
            candidate_skills=["Python", "Machine Learning", "Docker"],
            target_required_skills=["Python", "Machine Learning", "Docker", "Kubernetes"],
            resume_text="Experienced Python Machine Learning Developer skilled in Docker.",
            job_description_text="Looking for a Python Machine Learning developer with Docker experience.",
            years_experience=2.0,
            required_years=2.0,
            has_degree_match=True,
            document_layout_quality=90.0
        )
        
        self.assertEqual(factors.required_skill_match, 75.0)
        overall = compute_weighted_ats_overall(factors)
        self.assertTrue(0.0 <= overall <= 100.0)

    def test_resume_parser_and_doc_service(self):
        doc_service = DocumentService()
        text_sample = """
        John Doe
        john.doe@example.com | +1 555-0199
        
        Professional Summary:
        Machine Learning Engineer with 2 years of experience in Python, TensorFlow, and SQL.
        
        Education:
        B.Tech in Computer Science from Tech University
        
        Projects:
        BERT Sentiment Analysis System - Developed NLP sentiment classifier using PyTorch.
        """
        
        doc_analysis = doc_service.analyze_layout(b"pdf content", "resume.pdf", text_sample)
        self.assertGreaterEqual(doc_analysis.page_count, 1)

        parser = ResumeParser()
        profile = parser.parse_resume(text_sample, doc_analysis)
        
        self.assertEqual(profile.personal_information.name, "John Doe")
        self.assertEqual(profile.personal_information.email, "john.doe@example.com")
        self.assertIn("Python", profile.technical_skills)

    def test_ats_and_skill_gap_services(self):
        p = CandidateProfile(
            personal_information=PersonalInfo(name="Alice", email="alice@test.com"),
            technical_skills=["Python", "Machine Learning", "SQL"],
            projects=[Project(title="ML Pipeline", description="Built ML pipeline in Python.")]
        )
        
        ats_svc = ATSService()
        ats_res = ats_svc.analyze_ats_compatibility(p, "Machine Learning Engineer", "Requires Python, Machine Learning, Docker, Kubernetes")
        self.assertGreater(ats_res.overall_score, 0)
        self.assertIn("Docker", ats_res.required_skills_missing)

        sg_svc = SkillGapService()
        sg_res = sg_svc.analyze_skill_gaps(p, "Machine Learning Engineer", "Requires Python, Machine Learning, Docker, Kubernetes")
        self.assertIn("Python", sg_res.strong_match)
        self.assertIn("Docker", sg_res.missing_skills)
        self.assertGreater(len(sg_res.roadmap), 0)

        plan = sg_svc.generate_30_day_action_plan(sg_res, "Machine Learning Engineer")
        self.assertIn("30-Day Strategy", plan["summary_digest"])
        self.assertIn("Days 1–10", plan["phase1"]["title"])

    def test_interview_service(self):
        p = CandidateProfile(
            personal_information=PersonalInfo(name="Alice"),
            technical_skills=["Python", "TensorFlow"],
            projects=[Project(title="BERT Model", description="Trained BERT classifier.")]
        )
        int_svc = InterviewService()
        questions = int_svc.generate_personalized_questions(p, "Machine Learning Engineer")
        self.assertEqual(len(questions), 6)

        eval_res = int_svc.evaluate_answer(
            questions[0],
            "I chose BERT because of its bidirectional self-attention transformer architecture, which captures contextual semantics better than static embeddings."
        )
        self.assertGreater(eval_res.overall_answer_score, 70.0)

        report = int_svc.generate_session_report([eval_res])
        self.assertGreater(report["overall_score"], 0)
        self.assertIn("Interview Ready", report["readiness_level"])

    def test_career_and_profile_services(self):
        p = CandidateProfile(
            personal_information=PersonalInfo(name="Bob", email="bob@test.com"),
            technical_skills=["Python", "Machine Learning", "Docker"],
            education=[Education(degree="B.S. in Data Science", institution="State University")],
            projects=[Project(title="AI App", description="Streamlit AI dashboard.")]
        )
        
        career_svc = CareerService()
        roles = career_svc.recommend_career_roles(p)
        self.assertGreater(len(roles), 0)

        readiness = career_svc.evaluate_industry_readiness(p, ats_score=80.0)
        self.assertGreater(readiness.overall_score, 0)

        rec_svc = RecommendationService()
        projs = rec_svc.generate_project_recommendations(p, "ML Engineer", ["MLOps", "Kubernetes"])
        self.assertGreater(len(projs), 0)

        prof_svc = ProfileService()
        prof = prof_svc.generate_professional_profile(p, "ML Engineer")
        self.assertTrue("Bob" in prof.linkedin_headline or "ML Engineer" in prof.linkedin_headline)

    def test_database_operations(self):
        init_db()
        c_id = save_candidate_profile(
            filename="test_resume.pdf",
            file_type="Digital PDF",
            raw_text="Sample raw text",
            profile_dict={"personal_information": {"name": "Test User"}},
            doc_analysis_dict={"page_count": 1}
        )
        self.assertIsNotNone(c_id)
        self.assertGreater(c_id, 0)


if __name__ == "__main__":
    unittest.main()
